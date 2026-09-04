from io import BytesIO

import pytest

from specguard.documents import DocumentExtractionError, extract_text


def test_docx_table_stays_between_paragraphs():
    from docx import Document

    document = Document()
    document.add_paragraph("Before")
    table = document.add_table(rows=2, cols=2)
    for row, values in zip(table.rows, [("Field", "Type"), ("id", "string")]):
        for cell, value in zip(row.cells, values):
            cell.text = value
    document.add_paragraph("After")
    stream = BytesIO()
    document.save(stream)
    text = extract_text("test.docx", stream.getvalue(), max_chars=1000)
    assert text.index("Before") < text.index("| id | string |") < text.index("After")


def test_pdf_grid_preserves_cells_and_surrounding_text():
    from pypdf import PdfWriter
    from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

    writer = PdfWriter()
    page = writer.add_blank_page(width=400, height=400)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})}
    )
    stream = DecodedStreamObject()
    stream.set_data(
        b"50 200 300 100 re S 200 200 m 200 300 l S 50 250 m 350 250 l S "
        b"BT /F1 12 Tf 50 340 Td (Before) Tj ET "
        b"BT /F1 12 Tf 60 270 Td (Field) Tj ET "
        b"BT /F1 12 Tf 210 270 Td (Type) Tj ET "
        b"BT /F1 12 Tf 60 220 Td (id) Tj ET "
        b"BT /F1 12 Tf 210 220 Td (string) Tj ET "
        b"BT /F1 12 Tf 50 150 Td (After) Tj ET"
    )
    page[NameObject("/Contents")] = writer._add_object(stream)
    output = BytesIO()
    writer.write(output)
    text = extract_text("table.pdf", output.getvalue(), max_chars=1000)
    assert "| Field | Type |" in text
    assert text.index("Before") < text.index("| id | string |") < text.index("After")
    assert text.count("string") == 1


def test_extract_text_file() -> None:
    text = extract_text("spec.txt", "Техническое задание".encode(), max_chars=100)
    assert text == "Техническое задание"


def test_reject_empty_document() -> None:
    with pytest.raises(DocumentExtractionError, match="не найден текст"):
        extract_text("spec.md", b"   ", max_chars=100)


def test_reject_unsupported_type() -> None:
    with pytest.raises(DocumentExtractionError, match="не поддерживается"):
        extract_text("spec.xlsx", b"data", max_chars=100)
